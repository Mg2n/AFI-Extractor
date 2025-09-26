# ========= CONFIG =========
from pathlib import Path
from docx import Document
from openpyxl import Workbook, load_workbook
import pdfplumber
import re, unicodedata

HERE       = Path(__file__).parent.resolve()
INPUT_DIR  = HERE
XLSX_PATH  = HERE / "All_AFIs.xlsx"
SHEET_NAME = "Sheet1"

# ========= NORMALIZATION =========
DASH_CHARS = "\u2010\u2011\u2012\u2013\u2014\u2015\u2212-"
ZERO_WIDTH = "\u200b\u200c\u200d\u2060\uFEFF\u200e\u200f"
NBSP       = "\u00A0"
DASH_RE    = re.compile(f"[{re.escape(DASH_CHARS)}]")

def norm(s: str) -> str:
    s = unicodedata.normalize("NFKC", s)
    for ch in ZERO_WIDTH: s = s.replace(ch, "")
    s = s.replace(NBSP, " ").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = DASH_RE.sub("-", s)
    return s.strip()

def tidy(s: str) -> str:
    s = re.sub(r"\s{2,}", " ", s)
    return s.strip(" -.()[]:;").strip()

# ========= HEADERS / ITEMS =========
HEADERS = ["AFI", "Classification", "Recommendation", "Entity", "EE/FA", "Source File"]
RE_AFI_HDR   = re.compile(r"(?i)^\s*areas?\s+(for|of)\s+improvement\s*:?\s*$")
RE_RECO_HDR  = re.compile(r"(?i)^\s*recommendations?\s*:?\s*$")
RE_PROCESS   = re.compile(r"(?i)^\s*process\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)*)\s+(.+)$")
RE_PROC_SIMPLE = re.compile(r"(?i)^\s*(Value|Operational|Business)\b(?:\s*[:\-–—]\s*(.+))?$")
RE_NUM_ITEM  = re.compile(r"^\s*(\d+)\s*-\s*(.+?)\s*$")

# ========= TOC DETECTOR =========
DOT_LEADER = re.compile(r".{2,}\.{3,}\s*\d+\s*$")
TOC_TITLE  = re.compile(r"(?i)\btable of contents\b|\bcontents\b")

def is_toc_page(lines):
    if any(TOC_TITLE.search(x) for x in lines):
        return True
    dots = sum(1 for x in lines if DOT_LEADER.search(x))
    return dots >= 3

# ========= Classification/Entity extraction =========
def extract_ce_anywhere(text: str):
    s = text
    m_last = None
    for m in re.finditer(r"\(([^)]*?)\)", s):
        m_last = m
    if m_last:
        inside = DASH_RE.sub("-", m_last.group(1)).strip()
        parts = inside.split("-", 1)
        cls   = tidy(parts[0]) if parts else ""
        ent   = tidy(parts[1]) if len(parts) > 1 else ""
        if cls or ent:
            s = (s[:m_last.start()] + " " + s[m_last.end():]).strip()
            return tidy(s), cls, ent
    m_kw = None
    for m in re.finditer(r"(?i)\b(major|other)\b", s):
        m_kw = m
    if m_kw:
        dash = s.find("-", m_kw.start())
        if dash != -1:
            cls = m_kw.group(1).capitalize()
            ent = tidy(s[dash+1:])
            if ent:
                s = tidy(s[:m_kw.start()])
                return s, cls, ent
    return tidy(s), "", ""

def extract_ce_across_lines(line_list, idx, first_line):
    line = first_line
    if "(" in line and ")" not in line:
        buf = [line[line.index("("):]]
        j = idx + 1
        while j < len(line_list):
            buf.append(line_list[j])
            if ")" in line_list[j]:
                break
            j += 1
        combined = norm(" ".join(buf))
        m_last = None
        for m in re.finditer(r"\(([^)]*?)\)", combined):
            m_last = m
        if m_last:
            inside = DASH_RE.sub("-", m_last.group(1)).strip()
            parts = inside.split("-", 1)
            cls   = tidy(parts[0]) if parts else ""
            ent   = tidy(parts[1]) if len(parts) > 1 else ""
            cleaned = norm(first_line.replace(combined[m_last.start():m_last.end()], " "))
            return tidy(cleaned), cls, ent, min(j, len(line_list)-1)
        return tidy(first_line), "", "", min(j, len(line_list)-1)
    clean, cls, ent = extract_ce_anywhere(first_line)
    return clean, cls, ent, idx

# ========= READERS =========
def yield_lines_docx(doc_path: Path):
    doc = Document(doc_path)
    for p in doc.paragraphs:
        t = norm(p.text)
        if t: yield t
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for raw in cell.text.splitlines():
                    t = norm(raw)
                    if t: yield t

def yield_lines_pdf(pdf_path: Path):
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines = [norm(x) for x in text.splitlines() if norm(x)]
            if is_toc_page(lines):
                continue
            for ln in lines:
                yield ln

def yield_lines_any(path: Path):
    if path.suffix.lower() == ".docx":
        yield from yield_lines_docx(path)
    elif path.suffix.lower() == ".pdf":
        yield from yield_lines_pdf(path)

# ========= EXCEL HELPERS =========
def open_or_create_workbook(path: Path, sheet: str):
    if path.exists():
        wb = load_workbook(path)
        ws = wb[sheet] if sheet in wb.sheetnames else wb.active
        if ws.max_row < 1:
            ws.append(HEADERS)
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = sheet
        ws.append(HEADERS)
    return wb, ws

def detect_columns(ws):
    names = {}
    for col in range(1, ws.max_column + 1):
        v = (ws.cell(1, col).value or "").strip().lower()
        if v: names[v] = col
    col_afi     = names.get("afi", 1)
    col_class   = names.get("classification", 2)
    col_reco    = names.get("recommendation", 3)
    col_entity  = names.get("entity", 4)
    col_process = names.get("ee/fa", 5)
    col_file    = names.get("source file", 7)
    return col_afi, col_class, col_reco, col_entity, col_process, col_file

def next_row(ws): 
    return max(ws.max_row + 1, 2)

def write_row(ws, row, cols, afi_text, cls, ent, reco_text, process_label, src_file):
    col_afi, col_class, col_reco, col_entity, col_process, col_file = cols
    clean_afi, c2, e2 = extract_ce_anywhere(afi_text)
    if not cls and c2: cls = c2
    if not ent and e2: ent = e2
    ws.cell(row, col_afi).value     = clean_afi
    ws.cell(row, col_class).value   = cls
    ws.cell(row, col_reco).value    = reco_text
    ws.cell(row, col_entity).value  = ent
    ws.cell(row, col_process).value = process_label
    ws.cell(row, col_file).value    = src_file

# ========= CORE PARSER =========
def process_file(path: Path, ws, cols, start_row: int) -> int:
    row = start_row
    in_afi = in_reco = False
    process_label = ""
    afis = []
    recs = {}
    cur_rec_num = None
    cur_rec_parts = []
    last_afi_idx = None

    def flush_reco():
        nonlocal cur_rec_num, cur_rec_parts
        if cur_rec_num is not None:
            txt = " | ".join([t for t in cur_rec_parts if t])
            if txt:
                prev = recs.get(cur_rec_num, "")
                recs[cur_rec_num] = prev + (" | " if prev else "") + txt
        cur_rec_num = None
        cur_rec_parts = []

    lines = list(yield_lines_any(path))
    i = 0
    while i < len(lines):
        line = lines[i]
        m_proc = RE_PROCESS.match(line)
        m_simple = RE_PROC_SIMPLE.match(line)
        if m_proc or m_simple:
            flush_reco()
            if afis or recs:
                seq = 1
                for a in afis:
                    if a["num"] is None:
                        a["num"] = seq; seq += 1
                for a in sorted(afis, key=lambda x: int(x["num"])):
                    n = str(a["num"]); rec = recs.get(n, "")
                    write_row(ws, row, cols, a["text"], a["cls"], a["ent"], rec, process_label, path.name)
                    row += 1
            in_afi = in_reco = False
            if m_proc:
                process_label = f"{m_proc.group(1)} – {m_proc.group(2)}"
            else:
                head = m_simple.group(1).capitalize()
                tail = m_simple.group(2) or ""
                process_label = f"{head}{(' – ' + tail) if tail else ''}"
            afis, recs = [], {}
            last_afi_idx = None
            i += 1
            continue

        if RE_AFI_HDR.match(line):
            flush_reco()
            in_afi, in_reco = True, False
            last_afi_idx = None
            i += 1
            continue
        if RE_RECO_HDR.match(line):
            flush_reco()
            in_afi, in_reco = False, True
            cur_rec_num = None
            cur_rec_parts = []
            i += 1
            continue

        if in_afi:
            if re.fullmatch(r"\([^)]*\)", line):
                _, c, e = extract_ce_anywhere(line)
                if last_afi_idx is not None and (c or e):
                    if not afis[last_afi_idx]["cls"]: afis[last_afi_idx]["cls"] = c
                    if not afis[last_afi_idx]["ent"]: afis[last_afi_idx]["ent"] = e
                i += 1
                continue

            m = RE_NUM_ITEM.match(line)
            if m:
                num  = int(m.group(1))
                body = m.group(2)
                clean, cls, ent, j = extract_ce_across_lines(lines, i, body)
                afis.append({"num": num, "text": clean, "cls": cls, "ent": ent})
                last_afi_idx = len(afis) - 1
                i = max(i, j) + 1
                continue

            if "(" in line:
                clean, cls, ent, j = extract_ce_across_lines(lines, i, line)
                if cls or ent:
                    afis.append({"num": None, "text": clean, "cls": cls, "ent": ent})
                    last_afi_idx = len(afis) - 1
                    i = j + 1
                    continue

            i += 1
            continue

        if in_reco:
            m = RE_NUM_ITEM.match(line)
            if m:
                flush_reco()
                cur_rec_num  = m.group(1)
                cur_rec_parts = [tidy(m.group(2))]
            else:
                if cur_rec_num is None:
                    cur_rec_num = "1"
                cur_rec_parts.append(tidy(line))
            i += 1
            continue

        i += 1

    flush_reco()
    if afis or recs:
        seq = 1
        for a in afis:
            if a["num"] is None:
                a["num"] = seq; seq += 1
        for a in sorted(afis, key=lambda x: int(x["num"])):
            n = str(a["num"]); rec = recs.get(n, "")
            write_row(ws, row, cols, a["text"], a["cls"], a["ent"], rec, process_label, path.name)
            row += 1

    return row

# ========= MAIN =========
def main():
    wb, ws = open_or_create_workbook(XLSX_PATH, SHEET_NAME)
    cols = detect_columns(ws)
    row  = next_row(ws)
    files = [p for p in INPUT_DIR.iterdir() if p.suffix.lower() in (".docx", ".pdf")]
    files = [p for p in files if not p.name.startswith("~$")]
    files.sort(key=lambda p: p.name.lower())
    print(f"Found {len(files)} files (.docx/.pdf).")
    for i, f in enumerate(files, 1):
        print(f"[{i}/{len(files)}] {f.name}")
        row = process_file(f, ws, cols, row)
    wb.save(XLSX_PATH)
    print(f"✅ Done → {XLSX_PATH}")

if __name__ == "__main__":
    main()
